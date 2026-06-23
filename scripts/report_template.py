#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Cancer Buddy Report Template v2
用法: python report_template.py <report_data.json> <output.docx> [--type brief|detailed]
"""

import sys
import json
import argparse
from datetime import datetime
from docx import Document
from docx.shared import Pt, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


# ── UI string helper ────────────────────────────────────────────────────────────
# Section headings and labels come from report_data["ui"] (written by Phase 2 in
# the user's chosen language). Fallback values are Chinese so the template works
# even when "ui" is absent.
_UI_DEFAULTS = {
    "cover_brief_title":    "病例总结（简要版）",
    "cover_detailed_title": "病例总结（详细版）",
    "cover_subtitle":       "Cancer Buddy · 自动生成 · 仅供参考",
    "h_patient_id":  "患者标识",
    "h_condition":   "病情概要",
    "h_molecular":   "核心分子检测",
    "h_imaging":     "主要病灶分布",
    "h_labs":        "关键实验室指标",
    "h_treatment":   "治疗史",
    "h_gaps":        "待完善检查 / 建议补充记录",
    "h_pathway":     "当前治疗路径",
    "h_flags":       "待确认项",
    "s1": "患者基本信息",
    "s2": "病情概要",
    "s3": "分子检测与标志物",
    "s4": "影像学评估",
    "s5": "实验室指标摘要",
    "s6": "治疗史",
    "s7": "治疗路径总结",
    "sA": "建议补充记录",
    "sB": "待确认项",
    "sC": "信息来源索引",
    "kv_date":  "确诊时间",
    "kv_site":  "原发部位",
    "kv_hist":  "病理类型",
    "kv_diff":  "分化程度",
    "kv_stage": "临床分期",
    "kv_init":  "初诊 / 复发",
    "kv_mets":  "转移部位",
    "kv_stat":  "目前治疗状态",
    "tx_no_history": "暂无既往治疗记录。",
    "tx_current_status": "当前状态",
    "tx_efficacy": "疗效",
    "tx_stop_reason": "停药原因",
    "tx_toxicity": "毒副反应",
    "tx_line_names": "一线,二线,三线,四线,五线,六线",
    "tx_line_fallback": "第{n}线",
    "cover_brief_heading":    "病情简要总结",
    "cover_detailed_heading": "病情详细总结",
    "cover_detailed_eyebrow": "-- 病情详细总结",
    "cover_report_date":      "报告日期：{date}",
    "cover_disclaimer_short": "仅用于临床交流参考，不替代主治医师的判断",
    "labs_empty":      "暂无检验数据。",
    "molecular_empty": "暂无分子检测数据。",
    "imaging_empty":   "暂无影像学报告。",
    "trend_section_title": "关键指标趋势",
    "gaps_critical_title":    "(紧急) 对后续分析至关重要（建议尽快补充）",
    "gaps_recommended_title": "(建议) 有助于提升分析精准度",
    "gaps_covered_title":     "(已覆盖) 已充分覆盖",
    "flags_all_clear":   "所有提取字段已通过可疑值检查，无待确认项。",
    "flags_red_title":   "(待确认) 以下 {n} 项需在使用本报告做决策前确认：",
    "flags_yellow_title":"(建议核对) 以下 {n} 项建议核对（不影响报告生成）：",
    "flags_suggested_action": "建议：{action}",
    "action_cat_hospital": "现医院可补检",
    "action_cat_archive":  "需调阅历史档案",
    "action_cat_referral": "需转诊专项检查",
    "action_cat_unavail":  "组织标本不可及",
    "kv_separator": "：",
    "kv_pending_keywords": "未检测,未取得,Pending,待,—,Not obtained,not obtained,not tested,Not tested",
    "pw_current":     "当前较可能的路径",
    "pw_bridge":      "桥接",
    "pw_alternative": "备选试验路径",
    "chart_snapshot_title": "关键检验指标快照",
    "chart_ref_range":   "参考范围",
    "chart_normal":      "正常",
    "chart_abnormal":    "异常",
    "trend_ref_range":   "参考区间",
    "trend_out_of_range":"超出范围",
    "trend_normal_range":"正常范围",
    "footer_page_prefix": "第 ",
    "footer_page_mid":    " 页 / 共 ",
    "footer_page_suffix": " 页",
    "trend_xlabel_date":  "日期",
    "meta_report_date": "报告日期",
    "meta_files_analyzed": "分析文件",
    "meta_files_unit": "{n} 份",
    "meta_language": "语言",
    "meta_language_value": "简体中文",
    "pi_name":      "患者姓名",
    "pi_sex_age":   "性别 / 年龄",
    "pi_hospital":  "就诊医院",
    "pi_doctor":    "主管医生",
    "pi_ecog":      "ECOG 体能评分",
    "pi_ecog_pending": "待医生评估",
    "pi_diagnosis": "临床诊断",
    "pi_admit_no":  "住院号",
    "pi_patient_id":"病员号",
    "pi_patient_code": "病历编号",
    "labs_th_date":     "日期",
    "labs_th_item":     "检验项目",
    "labs_th_category": "类别",
    "labs_th_result":   "结果",
    "labs_th_ref":      "参考值",
    "labs_th_meaning":  "临床意义",
    "mol_th_item":    "检测项目",
    "mol_th_status":  "结果 / 状态",
    "mol_th_priority":"优先级",
    "mol_th_meaning": "临床意义",
    "mol_priority_high":   "关键",
    "mol_priority_medium": "建议",
    "mol_priority_low":    "参考",
    "mol_missing_keywords": "未检测,未取得,Pending,待回报,Not tested,Not obtained,not tested,not obtained",
    "img_th_date_type": "日期 / 类型",
    "img_th_summary":   "影像学摘要",
    "src_th_module":    "模块",
    "src_th_datapoint": "数据点",
    "src_th_file":      "来源文件",
    "disclaimer": (
        "本报告由 Cancer Buddy 自动生成 | 生成时间：{gen} | "
        "分析文件数：{fn} | 待确认 {fr} · 建议核对 {fy} · 已通过 {fg} | "
        "本报告不替代主诊医生的临床判断，所有治疗决策须与医生确认。"
    ),
}
_UI = {}   # populated by build_brief / build_detailed from data["ui"]

def _u(key):
    """Return localized UI string: data["ui"][key] → fallback default.
    Uses explicit key-presence check (not `or`) so legitimate empty-string
    overrides (e.g. footer_page_suffix="" in English) are not silently
    replaced by the Chinese default."""
    if key in _UI:
        return _UI[key]
    return _UI_DEFAULTS.get(key, key)

# ── Design tokens ──────────────────────────────────────────────────────────────
PRIMARY   = RGBColor(0x1F, 0x3B, 0x5C)
ACCENT    = RGBColor(0x2E, 0x75, 0xB6)
ALERT_T   = RGBColor(0xC0, 0x00, 0x00)
WARN_T    = RGBColor(0xAD, 0x57, 0x00)
OK_T      = RGBColor(0x1B, 0x5E, 0x20)

HDR_BG   = "D6E4F0"
ALERT_BG = "FCE8E8"
WARN_BG  = "FFF3E0"

FONT_CN = "Microsoft YaHei"
FONT_EN = "Arial"

# A4 版心：2.5cm 边距 → 16cm = 9072 DXA
CONTENT_DXA = 9072

# 治疗线配色（左边框）
LINE_ACCENT_HEX = ["1A6EA8", "C0600A", "1A6E1A", "6A1A8A"]
LINE_BG_HEX     = ["EBF4FB", "FEF3E8", "EBF5EB", "F5EBF9"]


# ── Low-level helpers ──────────────────────────────────────────────────────────
def _set_font(run, size_pt, bold=False, italic=False, color=None):
    run.font.name = FONT_EN
    run.font.size = Pt(size_pt)
    run.font.bold = bold
    run.font.italic = italic
    if color:
        run.font.color.rgb = color
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.append(rFonts)
    rFonts.set(qn("w:eastAsia"), FONT_CN)


def _shading(cell, hex_color):
    tcPr = cell._tc.get_or_add_tcPr()
    for old in tcPr.findall(qn("w:shd")):
        tcPr.remove(old)
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def _borders(cell, color="D0D7DE", size=4):
    tcPr = cell._tc.get_or_add_tcPr()
    for old in tcPr.findall(qn("w:tcBorders")):
        tcPr.remove(old)
    tcB = OxmlElement("w:tcBorders")
    for side in ("top", "left", "bottom", "right"):
        el = OxmlElement(f"w:{side}")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), str(size))
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), color)
        tcB.append(el)
    tcPr.append(tcB)


def _left_border_only(cell, color, size=24):
    """左粗边框，其余无边框（治疗卡片用）。"""
    tcPr = cell._tc.get_or_add_tcPr()
    for old in tcPr.findall(qn("w:tcBorders")):
        tcPr.remove(old)
    tcB = OxmlElement("w:tcBorders")
    sides = {"top": ("single", 4, "E0E8F0"),
             "bottom": ("single", 4, "E0E8F0"),
             "right": ("single", 4, "E0E8F0"),
             "left": ("single", size, color)}
    for side, (val, sz, col) in sides.items():
        el = OxmlElement(f"w:{side}")
        el.set(qn("w:val"), val)
        el.set(qn("w:sz"), str(sz))
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), col)
        tcB.append(el)
    tcPr.append(tcB)


def _bottom_rule(para, color="2E75B6", size=8, space=4):
    pPr = para._element.get_or_add_pPr()
    for old in pPr.findall(qn("w:pBdr")):
        pPr.remove(old)
    pBdr = OxmlElement("w:pBdr")
    bot = OxmlElement("w:bottom")
    bot.set(qn("w:val"), "single")
    bot.set(qn("w:sz"), str(size))
    bot.set(qn("w:space"), str(space))
    bot.set(qn("w:color"), color)
    pBdr.append(bot)
    pPr.append(pBdr)


def _fix_table(table, col_widths_dxa):
    """固定列宽：tblGrid + tblLayout fixed + tcW。"""
    tbl = table._tbl
    tblPr = tbl.find(qn("w:tblPr"))
    if tblPr is None:
        tblPr = OxmlElement("w:tblPr")
        tbl.insert(0, tblPr)
    for tag in ("w:tblW", "w:tblLayout", "w:tblInd"):
        for old in tblPr.findall(qn(tag)):
            tblPr.remove(old)
    tblW = OxmlElement("w:tblW")
    tblW.set(qn("w:w"), str(sum(col_widths_dxa)))
    tblW.set(qn("w:type"), "dxa")
    tblPr.append(tblW)
    tblLayout = OxmlElement("w:tblLayout")
    tblLayout.set(qn("w:type"), "fixed")
    tblPr.append(tblLayout)
    old_grid = tbl.find(qn("w:tblGrid"))
    if old_grid is not None:
        tbl.remove(old_grid)
    tblGrid = OxmlElement("w:tblGrid")
    for w in col_widths_dxa:
        gc = OxmlElement("w:gridCol")
        gc.set(qn("w:w"), str(w))
        tblGrid.append(gc)
    tblPr.addnext(tblGrid)
    for row in table.rows:
        for ci, cell in enumerate(row.cells):
            if ci >= len(col_widths_dxa):
                continue
            tcPr = cell._tc.get_or_add_tcPr()
            for old in tcPr.findall(qn("w:tcW")):
                tcPr.remove(old)
            tcW = OxmlElement("w:tcW")
            tcW.set(qn("w:w"), str(col_widths_dxa[ci]))
            tcW.set(qn("w:type"), "dxa")
            tcPr.append(tcW)


def _page_setup(doc):
    for sec in doc.sections:
        sec.page_width    = Cm(21)
        sec.page_height   = Cm(29.7)
        sec.top_margin    = Cm(2.5)
        sec.bottom_margin = Cm(2.5)
        sec.left_margin   = Cm(2.5)
        sec.right_margin  = Cm(2.5)
    doc.styles["Normal"].paragraph_format.space_before = Pt(0)
    doc.styles["Normal"].paragraph_format.space_after  = Pt(3)


def _footer(doc):
    section = doc.sections[0]
    footer  = section.footer
    para    = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    para.clear()
    def _field(r, instr):
        fc1 = OxmlElement("w:fldChar"); fc1.set(qn("w:fldCharType"), "begin"); r._element.append(fc1)
        it  = OxmlElement("w:instrText"); it.text = instr; r._element.append(it)
        fc2 = OxmlElement("w:fldChar"); fc2.set(qn("w:fldCharType"), "end"); r._element.append(fc2)
    r1 = para.add_run(_u("footer_page_prefix")); _set_font(r1, 8, color=RGBColor(0x88, 0x88, 0x88)); _field(r1, " PAGE ")
    r2 = para.add_run(_u("footer_page_mid"));    _set_font(r2, 8, color=RGBColor(0x88, 0x88, 0x88)); _field(r2, " NUMPAGES ")
    r3 = para.add_run(_u("footer_page_suffix"));  _set_font(r3, 8, color=RGBColor(0x88, 0x88, 0x88))


# ── Section headings ───────────────────────────────────────────────────────────
def heading_brief(doc, text):
    """简要版：左蓝边栏 + 浅蓝底色。"""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after  = Pt(6)
    p.paragraph_format.left_indent  = Cm(0.4)
    pPr = p._element.get_or_add_pPr()
    # 浅蓝背景
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear"); shd.set(qn("w:color"), "auto"); shd.set(qn("w:fill"), "EBF3FB")
    for old in pPr.findall(qn("w:shd")): pPr.remove(old)
    pPr.append(shd)
    # 左边框
    pBdr = OxmlElement("w:pBdr")
    left = OxmlElement("w:left")
    left.set(qn("w:val"), "single"); left.set(qn("w:sz"), "28")
    left.set(qn("w:space"), "8");   left.set(qn("w:color"), "2E75B6")
    pBdr.append(left)
    for old in pPr.findall(qn("w:pBdr")): pPr.remove(old)
    pPr.append(pBdr)
    r = p.add_run(text)
    _set_font(r, 11, bold=True, color=PRIMARY)


def heading_numbered(doc, number, text):
    """详细版：[N] 蓝色徽章 + 标题文字 + 分隔线。"""
    BADGE = 544
    tbl = doc.add_table(rows=1, cols=2)
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
    _fix_table(tbl, [BADGE, CONTENT_DXA - BADGE])
    # 徽章
    bc = tbl.rows[0].cells[0]
    _shading(bc, "2E75B6")
    _borders(bc, color="2E75B6", size=4)
    bc.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    bp = bc.paragraphs[0]
    bp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    bp.paragraph_format.space_before = Pt(0)
    bp.paragraph_format.space_after  = Pt(0)
    _set_font(bp.add_run(str(number)), 10, bold=True, color=RGBColor(0xFF, 0xFF, 0xFF))
    # 标题
    tc = tbl.rows[0].cells[1]
    _borders(tc, color="2E75B6", size=4)
    tc.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    tp = tc.paragraphs[0]
    tp.paragraph_format.space_before = Pt(0)
    tp.paragraph_format.space_after  = Pt(0)
    tp.paragraph_format.left_indent  = Cm(0.4)
    _set_font(tp.add_run(text), 13, bold=True, color=PRIMARY)
    # 分隔线
    rule = doc.add_paragraph()
    _bottom_rule(rule, color="C5D8EC", size=4)
    rule.paragraph_format.space_before = Pt(0)
    rule.paragraph_format.space_after  = Pt(8)


# ── Covers ─────────────────────────────────────────────────────────────────────
def _short_name(patient, dx):
    """提取简短的疾病名称（括号前部分），用于封面副标题。"""
    raw = patient.get("diagnosis", "") or dx.get("primary_site", "")
    # 取第一个全角/半角括号之前的内容，去掉分期信息
    for sep in ("（", "("):
        if sep in raw:
            raw = raw.split(sep)[0]
    return raw.strip() or "病情总结"


def cover_brief(doc, data):
    patient = data.get("patient", {})
    dx      = data.get("diagnosis", {})
    # 大标题（固定文字）
    p = doc.add_paragraph()
    _set_font(p.add_run(_u("cover_brief_heading")), 22, bold=True, color=PRIMARY)
    p.paragraph_format.space_after = Pt(6)
    # 副标题：仅显示简短病名，不拼接长字段
    subtitle = _short_name(patient, dx)
    if subtitle:
        ps = doc.add_paragraph()
        _set_font(ps.add_run(subtitle), 11, color=RGBColor(0x44, 0x44, 0x44))
        ps.paragraph_format.space_after = Pt(4)
    # 日期 + 声明
    date = patient.get("report_date", data.get("generated_at","")[:10])
    pd = doc.add_paragraph()
    _set_font(pd.add_run(_u("cover_report_date").format(date=date)), 9, color=RGBColor(0x55, 0x55, 0x55))
    _set_font(pd.add_run("  |  " + _u("cover_disclaimer_short")), 9, color=RGBColor(0x88, 0x88, 0x88))
    pd.paragraph_format.space_after = Pt(6)
    # 分隔线
    rule = doc.add_paragraph()
    _bottom_rule(rule, color="2E75B6", size=8)
    rule.paragraph_format.space_after = Pt(12)


def cover_detailed(doc, data):
    patient = data.get("patient", {})
    dx      = data.get("diagnosis", {})
    # 面包屑
    pc = doc.add_paragraph()
    _set_font(pc.add_run(_u("cover_detailed_eyebrow")), 9, color=RGBColor(0x66, 0x66, 0x66))
    pc.paragraph_format.space_after = Pt(10)
    # 大标题（固定文字，参考005 PDF）
    pt = doc.add_paragraph()
    _set_font(pt.add_run(_u("cover_detailed_heading")), 26, bold=True, color=PRIMARY)
    pt.paragraph_format.space_after = Pt(8)
    # 副标题：仅显示简短病名，不拼接长字段
    subtitle = _short_name(patient, dx)
    if subtitle:
        ps = doc.add_paragraph()
        _set_font(ps.add_run(subtitle), 11, color=RGBColor(0x44, 0x44, 0x44))
        ps.paragraph_format.space_after = Pt(10)
    # 元数据格（报告日期 / 分析文件数 / 语言）
    date    = patient.get("report_date", data.get("generated_at","")[:10])
    files_n = data.get("files_analyzed", 0)
    meta    = [(_u("meta_report_date"), date),
               (_u("meta_files_analyzed"), _u("meta_files_unit").format(n=files_n)),
               (_u("meta_language"), _u("meta_language_value"))]
    MW = CONTENT_DXA // 4
    tbl = doc.add_table(rows=1, cols=4)
    _fix_table(tbl, [MW, MW, MW, CONTENT_DXA - 3*MW])
    for i, (lbl, val) in enumerate(meta):
        c = tbl.rows[0].cells[i]
        _borders(c, color="E8E8E8", size=2)
        p = c.paragraphs[0]
        p.paragraph_format.space_before = Pt(3)
        p.paragraph_format.space_after  = Pt(3)
        _set_font(p.add_run(lbl + "\n"), 8, color=RGBColor(0x88, 0x88, 0x88))
        _set_font(p.add_run(val), 10, bold=True, color=PRIMARY)
    _borders(tbl.rows[0].cells[3], color="E8E8E8", size=2)
    # 分隔线
    rule = doc.add_paragraph()
    _bottom_rule(rule, color="2E75B6", size=8)
    rule.paragraph_format.space_before = Pt(10)
    rule.paragraph_format.space_after  = Pt(14)


# ── Patient info table ─────────────────────────────────────────────────────────
def _info_cell(c, text, is_label):
    """填充信息表单元格。"""
    if is_label:
        _shading(c, "F2F6FB")
    _borders(c, color="D0DCEB", size=4)
    p = c.paragraphs[0]
    p.paragraph_format.space_before = Pt(5)
    p.paragraph_format.space_after  = Pt(5)
    p.paragraph_format.left_indent  = Cm(0.25)
    is_pending = (not is_label) and (str(text) in ("—", "未取得") or "待" in str(text))
    _set_font(p.add_run(str(text)), 9.5,
              bold=is_label,
              color=PRIMARY if is_label else (
                  RGBColor(0x88, 0x88, 0x88) if is_pending else RGBColor(0x1A, 0x1A, 0x1A)))


def patient_info_table(doc, patient, extended=False):
    """患者信息表。
    brief (extended=False): 4 列网格，参考 004 PDF 样式。
    detailed (extended=True): 2 列，逐行详细。
    """
    if extended:
        # ── detailed: 2-col 完整信息 ──────────────────────────────────
        LABEL_W = 2100
        VALUE_W = CONTENT_DXA - LABEL_W
        rows = [
            (_u("pi_name"),      patient.get("name", "—")),
            (_u("pi_sex_age"),   f"{patient.get('sex','—')} / {patient.get('age','—')}"),
            (_u("pi_hospital"),  patient.get("hospital", "—")),
            (_u("pi_doctor"),    patient.get("doctor", "—")),
            (_u("pi_ecog"),      patient.get("ecog", _u("pi_ecog_pending"))),
            (_u("pi_diagnosis"), patient.get("diagnosis", "—")),
            (_u("meta_report_date"), patient.get("report_date", "—")),
            (_u("pi_admit_no"),  patient.get("admission_no", "—")),
            (_u("pi_patient_id"),patient.get("patient_id", "—")),
            (_u("pi_patient_code"), patient.get("patient_code", "—")),
        ]
        tbl = doc.add_table(rows=len(rows), cols=2)
        tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
        _fix_table(tbl, [LABEL_W, VALUE_W])
        for ri, (lbl, val) in enumerate(rows):
            lc, vc = tbl.rows[ri].cells
            _info_cell(lc, lbl, is_label=True)
            _info_cell(vc, val, is_label=False)
    else:
        # ── brief: 4-col 网格，参考 004 PDF ──────────────────────────
        # 每半宽 = 4536 DXA；标签列 1600，值列 2936
        LW = 1600
        VW = CONTENT_DXA // 2 - LW  # 2936
        grid = [
            (_u("pi_sex_age"),   f"{patient.get('sex','—')} / {patient.get('age','—')}",
             _u("pi_hospital"),  patient.get("hospital", "—")),
            (_u("pi_ecog"),      patient.get("ecog", _u("pi_ecog_pending")),
             _u("pi_diagnosis"), patient.get("diagnosis", "—")),
            (_u("meta_report_date"), patient.get("report_date", "—"),
             _u("pi_patient_id"),    patient.get("patient_id", "—")),
        ]
        tbl = doc.add_table(rows=len(grid), cols=4)
        tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
        _fix_table(tbl, [LW, VW, LW, VW])
        for ri, (l1, v1, l2, v2) in enumerate(grid):
            cells = tbl.rows[ri].cells
            _info_cell(cells[0], l1, is_label=True)
            _info_cell(cells[1], v1, is_label=False)
            _info_cell(cells[2], l2, is_label=True)
            _info_cell(cells[3], v2, is_label=False)
    doc.add_paragraph()


# ── Labs ───────────────────────────────────────────────────────────────────────
def labs_cards(doc, labs, n=4):
    """卡片网格（简要版）。"""
    if not labs:
        _set_font(doc.add_paragraph(_u("labs_empty")).add_run(""), 10, color=RGBColor(0x88,0x88,0x88))
        return
    CW = CONTENT_DXA // n
    nrows = (len(labs) + n - 1) // n
    tbl = doc.add_table(rows=nrows, cols=n)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    _fix_table(tbl, [CW]*n)
    for idx, item in enumerate(labs):
        ri, ci = divmod(idx, n)
        c = tbl.rows[ri].cells[ci]
        flag = item.get("flag","normal")
        if flag in ("high","low"):
            _shading(c, ALERT_BG)
        _borders(c, color="D0D8E4", size=4)
        # 项目名
        p1 = c.paragraphs[0]
        p1.paragraph_format.space_before = Pt(6)
        p1.paragraph_format.space_after  = Pt(1)
        p1.paragraph_format.left_indent  = Cm(0.25)
        _set_font(p1.add_run(item.get("item","—")), 8, color=RGBColor(0x77,0x77,0x77))
        # 值
        p2 = c.add_paragraph()
        p2.paragraph_format.space_before = Pt(0)
        p2.paragraph_format.space_after  = Pt(1)
        p2.paragraph_format.left_indent  = Cm(0.25)
        val = item.get("value","—")
        if flag == "high": val += " +"
        elif flag == "low": val += " -"
        _set_font(p2.add_run(val), 10, bold=True,
                  color=ALERT_T if flag in ("high","low") else PRIMARY)
        # 参考区间
        ref = item.get("reference","")
        if ref:
            p3 = c.add_paragraph()
            p3.paragraph_format.space_before = Pt(0)
            p3.paragraph_format.space_after  = Pt(6)
            p3.paragraph_format.left_indent  = Cm(0.25)
            _set_font(p3.add_run(ref), 7.5, color=RGBColor(0xAA,0xAA,0xAA))
        else:
            c.add_paragraph().paragraph_format.space_after = Pt(6)
    # 补空格
    filled = len(labs) % n
    if filled:
        for ci in range(filled, n):
            _borders(tbl.rows[nrows-1].cells[ci], color="F0F0F0", size=2)
    doc.add_paragraph()


def labs_table(doc, labs):
    """标准表格（详细版）。列序：日期 → 检验项目 → 类别 → 结果 → 参考值 → 临床意义。"""
    if not labs:
        p = doc.add_paragraph(_u("labs_empty"))
        _set_font(p.runs[0] if p.runs else p.add_run(""), 10, color=RGBColor(0x88, 0x88, 0x88))
        return
    headers    = [_u("labs_th_date"), _u("labs_th_item"), _u("labs_th_category"),
                  _u("labs_th_result"), _u("labs_th_ref"), _u("labs_th_meaning")]
    col_widths = [1361,   1814,       998,    1270,   1179,    2450]   # sum = 9072
    tbl = doc.add_table(rows=1, cols=6)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    _fix_table(tbl, col_widths)
    for i, h in enumerate(headers):
        c = tbl.rows[0].cells[i]
        _shading(c, HDR_BG)
        _borders(c)
        p = c.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _set_font(p.add_run(h), 9.5, bold=True, color=PRIMARY)
    for item in labs:
        flag = item.get("flag", "normal")
        ab   = flag in ("high", "low")
        row  = tbl.add_row()
        vals = [item.get("date", "—"), item.get("item", "—"), item.get("category", "—"),
                item.get("value", "—"), item.get("reference", "—"), item.get("note", "—")]
        for i, val in enumerate(vals):
            c = row.cells[i]
            _borders(c)
            if ab:
                _shading(c, ALERT_BG)
            p = c.paragraphs[0]
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after  = Pt(2)
            if i == 3 and ab:
                arrow = " +" if flag == "high" else " -"
                _set_font(p.add_run(str(val) + arrow), 9.5, bold=True, color=ALERT_T)
            else:
                _set_font(p.add_run(str(val)), 9.5,
                          color=ALERT_T if (ab and i == 5) else RGBColor(0x1A, 0x1A, 0x1A))
    doc.add_paragraph()


# ── Lab trend helpers ──────────────────────────────────────────────────────────
def _parse_numeric(text):
    """从字符串提取第一个数值，返回 float 或 None。"""
    import re
    m = re.search(r'[-+]?\d+\.?\d*', str(text))
    return float(m.group()) if m else None


def _parse_ref_range(ref_str):
    """
    解析参考区间字符串，返回 (low, high)。
    支持：'3.5-5.0', '< 40', '≤ 40', '> 3', '≥ 3'；无法解析返回 (None, None)。
    """
    import re
    s = str(ref_str).strip()
    m = re.search(r'([\d.]+)\s*[~～\-\–—]\s*([\d.]+)', s)
    if m:
        a, b = float(m.group(1)), float(m.group(2))
        return (min(a, b), max(a, b))
    m = re.search(r'[<≤][=]?\s*([\d.]+)', s)
    if m:
        return (0.0, float(m.group(1)))
    m = re.search(r'[>≥][=]?\s*([\d.]+)', s)
    if m:
        return (float(m.group(1)), None)
    return (None, None)


def _mpl_cjk_setup():
    """尝试为 matplotlib 配置 CJK 字体，失败静默跳过。"""
    try:
        import matplotlib.font_manager as fm
        import matplotlib.pyplot as plt
        cjk = next((f for f in fm.findSystemFonts()
                    if any(x in f for x in ['NotoSansCJK', 'NotoSerifCJK', 'WenQuanYi'])),
                   None)
        if cjk:
            plt.rcParams['font.family'] = fm.FontProperties(fname=cjk).get_name()
        plt.rcParams['axes.unicode_minus'] = False
    except Exception:
        pass


def _trend_group_key(lab):
    """Compute base_item grouping key for trend charts / dedup logic."""
    import re as _re
    _SUFFIX = _re.compile(
        r'\s*——\s*.+$'
        r'|\s+(基线|复查|化疗前|化疗后|术后|随访|初诊|入院|出院).*$'
    )
    b = lab.get("base_item", "").strip()
    if b:
        return b
    return _SUFFIX.sub("", lab.get("item", "")).strip()


def _trend_keys(labs):
    """Return set of base_item group keys that have >=2 numeric measurements."""
    from collections import defaultdict
    counts = defaultdict(int)
    for lab in labs:
        if _parse_numeric(lab.get("value", "")) is not None and lab.get("date", "").strip():
            counts[_trend_group_key(lab)] += 1
    return {k for k, c in counts.items() if c >= 2}


def _latest_per_trend(labs, trend_keys):
    """
    For items whose group key is in trend_keys (shown in trend chart):
      keep only the row with the latest date.
    All other items: keep as-is.
    Preserves original order (first occurrence wins the slot for trend items).
    """
    latest = {}
    for lab in labs:
        k = _trend_group_key(lab)
        if k in trend_keys:
            if k not in latest or lab.get("date","") > latest[k].get("date",""):
                latest[k] = lab

    result = []
    seen = set()
    for lab in labs:
        k = _trend_group_key(lab)
        if k in trend_keys:
            if k not in seen:
                result.append(latest[k])
                seen.add(k)
        else:
            result.append(lab)
    return result


def _date_to_xpos(ev_date, xs_dates):
    """
    Linearly interpolate ev_date to a float x-index within [0, n-1].
    xs_dates: list of datetime objects, sorted ascending.
    Returns None if outside range.
    """
    n = len(xs_dates)
    if not n:
        return None
    if ev_date < xs_dates[0] or ev_date > xs_dates[-1]:
        return None
    if ev_date == xs_dates[0]:
        return 0.0
    if ev_date == xs_dates[-1]:
        return float(n - 1)
    for i in range(n - 1):
        if xs_dates[i] <= ev_date <= xs_dates[i + 1]:
            total = (xs_dates[i + 1] - xs_dates[i]).total_seconds()
            frac  = (ev_date - xs_dates[i]).total_seconds() / total if total else 0.0
            return i + frac
    return None


def labs_snapshot_chart(doc, labs, png_save_path=None):
    """
    单时间点快照图：显示所有有参考范围的数值指标当前值与参考区间的对比。
    单时间点和多时间点均可调用（多时间点时显示最新值）。
    doc: Document 对象（插入图片用）；若 None 则只保存到 png_save_path。
    png_save_path: 若指定，同时保存 PNG 文件（供 .md 引用）。
    返回 png_save_path（保存成功）或 None。
    """
    try:
        import matplotlib
        matplotlib.use('Agg')
        import matplotlib.pyplot as plt
    except ImportError:
        return None
    import tempfile, os, re as _re

    _mpl_cjk_setup()

    # ── 收集可绘制项 ──────────────────────────────────────────────────────────
    # 每个 base_item 只取最新一条
    latest = {}
    for lab in labs:
        raw_val = lab.get("value", "")
        # 跳过批量项（含 3 个以上 "/" 分隔的子项）
        if raw_val.count("/") >= 2:
            continue
        val = _parse_numeric(raw_val)
        if val is None:
            continue
        ref_low, ref_high = _parse_ref_range(lab.get("reference", ""))
        if ref_low is None and ref_high is None:
            continue
        key = _trend_group_key(lab)
        date = lab.get("date", "")
        if key not in latest or date > latest[key]["date"]:
            latest[key] = {
                "label": (lab.get("base_item") or lab.get("item", ""))[:22],
                "value": val,
                "flag": lab.get("flag", "normal"),
                "ref_low": ref_low if ref_low is not None else 0.0,
                "ref_high": ref_high,
                "date": date,
            }

    items = list(latest.values())
    if not items:
        return None

    # ── 绘图 ──────────────────────────────────────────────────────────────────
    n = len(items)
    row_h = 0.52
    fig_h = max(2.4, n * row_h + 1.0)
    fig, ax = plt.subplots(figsize=(7.5, fig_h))
    fig.patch.set_facecolor('white')
    ax.set_facecolor('white')

    y_positions = list(range(n - 1, -1, -1))   # top→bottom

    for idx, (item, ypos) in enumerate(zip(items, y_positions)):
        val = item["value"]
        rlo = item["ref_low"]
        rhi = item["ref_high"]
        flag = item["flag"]
        dot_color = '#dc3545' if flag in ('high', 'low', 'critical') else '#198754'

        # Compute x-axis span for this item (normalize within the row strip)
        lo = 0.0
        hi = (rhi * 1.6) if rhi else (val * 2.2 if val > 0 else 10.0)
        hi = max(hi, val * 1.2)

        # Normalize to [0, 1] data coords within item's strip (height = 0.7)
        def norm(x):
            return (x - lo) / (hi - lo) if hi > lo else 0.5

        strip_y0 = ypos - 0.35
        strip_y1 = ypos + 0.35

        # Reference band
        if rhi:
            x0_ref = norm(rlo)
            x1_ref = norm(rhi)
            ax.barh(ypos, x1_ref - x0_ref, left=x0_ref, height=0.58,
                    color='#d4edda', edgecolor='#6c9e77', linewidth=0.6,
                    zorder=2, align='center')

        # Baseline
        ax.axhline(ypos, color='#dee2e6', lw=0.5, zorder=1)

        # Value dot
        xv = norm(val)
        ax.scatter([xv], [ypos], s=90, color=dot_color, zorder=5, linewidths=0)

        # Stem from ref band right edge to dot (if high)
        if rhi and val > rhi:
            ax.plot([norm(rhi), xv], [ypos, ypos], color=dot_color, lw=1.2, zorder=4)
        elif rlo > 0 and val < rlo:
            ax.plot([xv, norm(rlo)], [ypos, ypos], color=dot_color, lw=1.2, zorder=4)

        # Value label
        arrow = "↑" if flag == 'high' else ("↓" if flag == 'low' else "")
        label_x = min(xv + 0.03, 0.97)
        ax.text(label_x, ypos + 0.22, f"{val}{arrow}",
                ha='left', va='bottom', fontsize=7.5,
                color=dot_color, fontweight='bold', zorder=6)

    # Y-axis labels
    ax.set_yticks(y_positions)
    ax.set_yticklabels([it["label"] for it in items], fontsize=8.5)
    ax.set_xlim(0, 1)
    ax.set_ylim(-0.6, n - 0.4)
    ax.set_xticks([])
    for spine in ['top', 'right', 'bottom']:
        ax.spines[spine].set_visible(False)
    ax.spines['left'].set_linewidth(0.5)

    # Legend
    import matplotlib.patches as mpatches
    ax.legend(
        handles=[
            mpatches.Patch(color='#d4edda', edgecolor='#6c9e77', label=_u('chart_ref_range')),
            plt.scatter([], [], s=60, color='#198754', label=_u('chart_normal')),
            plt.scatter([], [], s=60, color='#dc3545', label=_u('chart_abnormal')),
        ],
        loc='lower right', fontsize=7.5, framealpha=0.85, edgecolor='#dee2e6'
    )

    ax.set_title(_u("chart_snapshot_title"), fontsize=10, fontweight='bold', pad=6)
    plt.tight_layout()

    # Save / embed
    tmp = None
    try:
        if png_save_path:
            os.makedirs(os.path.dirname(png_save_path), exist_ok=True)
            plt.savefig(png_save_path, dpi=150, bbox_inches='tight', facecolor='white')
        else:
            tmp = tempfile.NamedTemporaryFile(suffix=".png", delete=False)
            tmp.close()
            png_save_path = tmp.name
            plt.savefig(png_save_path, dpi=150, bbox_inches='tight', facecolor='white')
        plt.close()
        if doc:
            from docx.shared import Inches
            doc.add_picture(png_save_path, width=Inches(5.8))
        return png_save_path if not tmp else None
    except Exception as e:
        print(f"[snapshot] chart error: {e}")
        plt.close()
        return None
    finally:
        if tmp and os.path.exists(tmp.name):
            try:
                os.unlink(tmp.name)
            except Exception:
                pass


def labs_trend_charts(doc, labs, trend_events=None):
    """
    关键指标趋势图：每行 2 列子图，每张图最多 2 行（4 个指标），防止跨页。
    - 优先使用 base_item 字段分组；fallback 到正则规范化
    - 参考区间 = 蓝色背景带 + 虚线上下限
    - 超出参考范围的点/值标签 = 红色；正常范围内 = 蓝色
    - trend_events: [{date, label}] — 在落入该指标 X 轴范围内的事件画竖虚线，
      并将事件标签写入该子图自己的图例（右上角），帮助读者理解趋势转变原因
    - X 轴格式：同年 -> MM-DD；跨年 -> YYYY-MM-DD；点多(>8) -> YYYY-MM
    - 每张图独立保存并插入 doc，避免整体过高导致跨页
    """
    try:
        import matplotlib
        matplotlib.use('Agg')
        import matplotlib.pyplot as plt
        import matplotlib.patches as mpatches
    except ImportError:
        return
    import tempfile, os
    from datetime import datetime
    from collections import defaultdict, OrderedDict

    _mpl_cjk_setup()

    DATE_FMTS = ["%Y-%m-%d", "%Y/%m/%d", "%Y%m%d", "%Y年%m月%d日"]
    trend_events = trend_events or []

    # -- Group by base_item ------------------------------------------------
    by_item = defaultdict(list)
    for lab in labs:
        val  = _parse_numeric(lab.get("value", ""))
        date = lab.get("date", "").strip()
        if val is not None and date:
            by_item[_trend_group_key(lab)].append(lab)

    multi = OrderedDict(
        (k, sorted(v, key=lambda x: x.get("date", "")))
        for k, v in by_item.items() if len(v) >= 2
    )
    if not multi:
        return

    # -- Parse time series -------------------------------------------------
    series = {}
    for key, records in multi.items():
        ref_str   = next((r.get("reference","") for r in records if r.get("reference")), "")
        ref_low, ref_high = _parse_ref_range(ref_str)
        pts = []
        for r in records:
            v  = _parse_numeric(r.get("value", ""))
            fl = r.get("flag", "normal")
            d  = None
            for fmt in DATE_FMTS:
                try:
                    d = datetime.strptime(r.get("date",""), fmt); break
                except ValueError:
                    pass
            if d and v is not None:
                pts.append((d, v, fl))
        if len(pts) < 2:
            continue
        pts.sort(key=lambda x: x[0])
        xs, ys, fs = zip(*pts)
        series[key] = dict(xs=xs, ys=ys, fs=fs,
                           ref_low=ref_low, ref_high=ref_high, ref_str=ref_str)

    if not series:
        return

    # -- Section sub-heading -----------------------------------------------
    ph = doc.add_paragraph()
    ph.paragraph_format.space_before = Pt(10)
    ph.paragraph_format.space_after  = Pt(4)
    _set_font(ph.add_run(_u("trend_section_title")), 10, bold=True, color=PRIMARY)
    _bottom_rule(ph, color="C5D8EC", size=4, space=3)

    # -- Parse clinical event dates ----------------------------------------
    parsed_events = []
    for ev in trend_events:
        for fmt in DATE_FMTS:
            try:
                parsed_events.append({
                    "dt":    datetime.strptime(ev.get("date",""), fmt),
                    "label": ev.get("label",""),
                })
                break
            except ValueError:
                pass

    # -- Layout constants --------------------------------------------------
    NCOLS           = 2
    MAX_ROWS_PER_FIG = 2     # max rows per figure -> max 4 items -> fits on one A4 page
    CHUNK_SIZE      = MAX_ROWS_PER_FIG * NCOLS
    FIG_W           = 6.0   # inches ~15 cm
    ROW_H           = 2.85  # inches per row
    LEG_H           = 0.50  # inches reserved for figure-level legend + x-labels

    items_list = list(series.items())
    chunks = [items_list[i:i+CHUNK_SIZE] for i in range(0, len(items_list), CHUNK_SIZE)]

    # -- Shared figure-level legend handles (same for every chunk) ---------
    fig_legend_handles = [
        mpatches.Patch(facecolor="#2E75B6", alpha=0.3, label=_u("trend_ref_range")),
        plt.Line2D([0],[0], marker="o", color="w",
                   markerfacecolor="#C00000", markersize=6, label=_u("trend_out_of_range")),
        plt.Line2D([0],[0], marker="o", color="w",
                   markerfacecolor="#2E75B6", markersize=6, label=_u("trend_normal_range")),
    ]

    for chunk in chunks:
        n_items_chunk = len(chunk)
        NROWS = (n_items_chunk + NCOLS - 1) // NCOLS
        fig_h = ROW_H * NROWS + LEG_H

        fig, axes = plt.subplots(NROWS, NCOLS,
                                 figsize=(FIG_W, fig_h),
                                 dpi=150, squeeze=False)
        fig.patch.set_facecolor("#FFFFFF")

        for idx, (key, s) in enumerate(chunk):
            row, col = divmod(idx, NCOLS)
            ax = axes[row][col]
            ax.set_facecolor("#FAFBFC")

            xs, ys, fs = s["xs"], s["ys"], s["fs"]
            ref_low, ref_high, ref_str = s["ref_low"], s["ref_high"], s["ref_str"]

            # Y range
            all_v = list(ys)
            bands = [v for v in (ref_low, ref_high) if v is not None]
            combo = all_v + bands
            y_lo  = min(combo) * (1.15 if min(combo) < 0 else 0.85)
            y_hi  = max(combo) * 1.22
            y_lo  = min(y_lo, 0) if min(all_v) <= 0 else y_lo

            # Reference band
            if ref_low is not None and ref_high is not None:
                ax.axhspan(ref_low, ref_high, alpha=0.15, color="#2E75B6", zorder=1)
                ax.axhline(ref_low,  color="#2E75B6", lw=0.8, ls="--", alpha=0.5)
                ax.axhline(ref_high, color="#2E75B6", lw=0.8, ls="--", alpha=0.5)
            elif ref_high is not None:
                ax.axhspan(y_lo, ref_high, alpha=0.15, color="#2E75B6", zorder=1)
                ax.axhline(ref_high, color="#2E75B6", lw=0.8, ls="--", alpha=0.5)
            elif ref_low is not None:
                ax.axhspan(ref_low, y_hi, alpha=0.15, color="#2E75B6", zorder=1)
                ax.axhline(ref_low,  color="#2E75B6", lw=0.8, ls="--", alpha=0.5)

            ax.set_ylim(y_lo, y_hi)

            # Line + scatter
            n    = len(xs)
            xidx = list(range(n))
            ax.plot(xidx, ys, color="#BBBBBB", lw=1.2, zorder=2)
            for i, (y, fl) in enumerate(zip(ys, fs)):
                clr = "#C00000" if fl in ("high", "low") else "#2E75B6"
                ax.scatter(i, y, color=clr, s=40, zorder=4)
                ax.annotate(str(y), (i, y),
                            textcoords="offset points", xytext=(0, 6),
                            ha="center", fontsize=6.5, color=clr, fontweight="bold")

            # Clinical event annotations:
            # - draw vertical dotted line on the subplot
            # Clinical event annotations — treatment interventions only (surgery/chemo/regimen change).
            # Stagger labels vertically when events are too close on x-axis to prevent overlap.
            xs_list = list(xs)
            ax_event_handles = []
            prev_xpos = None
            stagger_level = 0   # alternates 0 / 1 for nearby events
            MIN_X_GAP = 0.25    # min x-distance before staggering kicks in
            Y_TOP   = y_hi * 0.97
            Y_MID   = y_hi * 0.72  # alternate label start height
            for ev in parsed_events:
                xpos = _date_to_xpos(ev["dt"], xs_list)
                if xpos is None:
                    continue
                ax.axvline(xpos, color="#E07000", lw=1.0, ls="--", alpha=0.75, zorder=3)
                # Choose label y-start based on proximity to previous event
                if prev_xpos is not None and abs(xpos - prev_xpos) < MIN_X_GAP:
                    stagger_level = 1 - stagger_level   # flip
                else:
                    stagger_level = 0
                label_y = Y_MID if stagger_level else Y_TOP
                ax.text(xpos + 0.06, label_y,
                        ev["label"], fontsize=5.5, color="#C05000",
                        ha="left", va="top", rotation=90, clip_on=True,
                        fontweight="bold")
                ax_event_handles.append(
                    plt.Line2D([0],[0], color="#E07000", lw=1.0, ls="--", alpha=0.8,
                               label=ev["label"])
                )
                prev_xpos = xpos

            # Subplot-level legend: only treatment events (compact, top-right)
            # This is the primary visual explanation of WHY the trend changed.
            if ax_event_handles:
                ax.legend(handles=ax_event_handles,
                          fontsize=5.5, loc="upper right",
                          title="治疗事件", title_fontsize=5.5,
                          framealpha=0.88, edgecolor="#DDDDDD",
                          borderaxespad=0.3, handlelength=1.5,
                          labelspacing=0.25)

            # Adaptive X-axis date format
            year_set = {d.year for d in xs}
            if n > 8:
                x_fmt = "%Y-%m"
            elif len(year_set) > 1:
                x_fmt = "%Y-%m-%d"
            else:
                x_fmt = "%m-%d"

            if n > 6:
                tick_idx    = xidx[::2]
                tick_labels = [xs[i].strftime(x_fmt) for i in range(0, n, 2)]
            else:
                tick_idx    = xidx
                tick_labels = [d.strftime(x_fmt) for d in xs]

            rot = 35 if n > 3 else 0
            ha  = "right" if n > 3 else "center"
            ax.set_xticks(tick_idx)
            ax.set_xticklabels(tick_labels, fontsize=6.5, rotation=rot, ha=ha)
            ax.tick_params(axis="x", pad=2)
            ax.tick_params(axis="y", labelsize=6.5)
            # X-axis label: clarify axis represents dates
            ax.set_xlabel(_u("trend_xlabel_date"), fontsize=7, color="#666666", labelpad=3)

            # Subplot title: indicator name + reference string
            title = key
            if ref_str and ref_str not in ("—", ""):
                title += "\n(ref: " + ref_str + ")"
            ax.set_title(title, fontsize=7.5, fontweight="bold",
                         color="#1F3B5C", pad=4, loc="left")
            ax.grid(axis="y", alpha=0.3, lw=0.5, color="#CCCCCC")
            for sp in ax.spines.values():
                sp.set_edgecolor("#DDDDDD")
            ax.set_xlim(-0.3, n - 0.7)

        # Hide unused subplots
        for idx in range(n_items_chunk, NROWS * NCOLS):
            row, col = divmod(idx, NCOLS)
            axes[row][col].set_visible(False)

        # Figure-level legend: shared symbols (参考区间 / 超出范围 / 正常范围)
        bottom_frac = LEG_H / fig_h
        plt.tight_layout(pad=0.7, h_pad=1.6, w_pad=1.0,
                         rect=[0, bottom_frac, 1, 1])
        fig.legend(handles=fig_legend_handles, fontsize=7, loc="lower center",
                   bbox_to_anchor=(0.5, 0.0),
                   framealpha=0.92, edgecolor="#DDDDDD",
                   ncol=3, borderaxespad=0.2)

        # Save + insert
        tf = tempfile.NamedTemporaryFile(suffix=".png", delete=False)
        tf.close()
        fig.savefig(tf.name, dpi=150, bbox_inches="tight",
                    facecolor="white", edgecolor="none")
        plt.close(fig)

        img_p = doc.add_paragraph()
        img_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        img_p.add_run().add_picture(tf.name, width=Cm(15))
        os.unlink(tf.name)

    doc.add_paragraph().paragraph_format.space_after = Pt(6)



# ── Molecular ──────────────────────────────────────────────────────────────────
def molecular_table(doc, molecular):
    if not molecular:
        _set_font(doc.add_paragraph(_u("molecular_empty")).add_run(""), 10, color=RGBColor(0x88,0x88,0x88))
        return
    headers   = [_u("mol_th_item"), _u("mol_th_status"), _u("mol_th_priority"), _u("mol_th_meaning")]
    col_widths = [2268, 2540, 998, 3266]
    tbl = doc.add_table(rows=1, cols=4)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    _fix_table(tbl, col_widths)
    for i, h in enumerate(headers):
        c = tbl.rows[0].cells[i]
        _shading(c, HDR_BG); _borders(c)
        p = c.paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _set_font(p.add_run(h), 9.5, bold=True, color=PRIMARY)
    PMAP = {"high": (ALERT_BG, ALERT_T, _u("mol_priority_high")),
            "medium": (WARN_BG, WARN_T, _u("mol_priority_medium")),
            "low": ("", OK_T, _u("mol_priority_low"))}
    for item in molecular:
        bg, tc, lbl = PMAP.get(item.get("priority","medium"), (WARN_BG, WARN_T, _u("mol_priority_medium")))
        status = item.get("status","—")
        missing = any(kw in status for kw in _u("mol_missing_keywords").split(","))
        row = tbl.add_row()
        for i, val in enumerate([item.get("item","—"), status, lbl, item.get("note","—")]):
            c = row.cells[i]; _borders(c)
            if bg: _shading(c, bg)
            p = c.paragraphs[0]
            p.paragraph_format.space_before = Pt(2); p.paragraph_format.space_after = Pt(2)
            if i==1 and missing:
                _set_font(p.add_run(str(val)), 9.5, color=RGBColor(0x88,0x88,0x88))
            elif i==2:
                _set_font(p.add_run(str(val)), 9.5, bold=True, color=tc)
            else:
                _set_font(p.add_run(str(val)), 9.5, color=RGBColor(0x1A,0x1A,0x1A))
    doc.add_paragraph()


# ── Imaging ────────────────────────────────────────────────────────────────────
def imaging_section(doc, imaging):
    note  = imaging.get("note")  if imaging else None
    items = imaging.get("items", []) if imaging else []
    if note:
        p = doc.add_paragraph(note)
        _set_font(p.runs[0] if p.runs else p.add_run(""), 10, color=RGBColor(0x88,0x88,0x88))
        doc.add_paragraph(); return
    if not items:
        p = doc.add_paragraph(_u("imaging_empty"))
        _set_font(p.runs[0] if p.runs else p.add_run(""), 10, color=RGBColor(0x88,0x88,0x88))
        doc.add_paragraph(); return
    col_widths = [2268, CONTENT_DXA - 2268]
    tbl = doc.add_table(rows=1, cols=2)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    _fix_table(tbl, col_widths)
    for i, h in enumerate([_u("img_th_date_type"), _u("img_th_summary")]):
        c = tbl.rows[0].cells[i]; _shading(c, HDR_BG); _borders(c)
        p = c.paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _set_font(p.add_run(h), 9.5, bold=True, color=PRIMARY)
    for item in items:
        row = tbl.add_row()
        for i, val in enumerate([item.get("date_type","—"), item.get("summary","—")]):
            c = row.cells[i]; _borders(c)
            p = c.paragraphs[0]
            p.paragraph_format.space_before = Pt(3); p.paragraph_format.space_after = Pt(3)
            _set_font(p.add_run(str(val)), 9.5,
                      bold=(i==0), color=PRIMARY if i==0 else RGBColor(0x1A,0x1A,0x1A))
    doc.add_paragraph()


# ── Treatment timeline ─────────────────────────────────────────────────────────
def treatment_history(doc, treatment):
    lines   = treatment.get("lines", [])
    current = treatment.get("current", "")
    note    = treatment.get("note", "")
    if note:
        p = doc.add_paragraph(note)
        _set_font(p.runs[0] if p.runs else p.add_run(""), 10, color=RGBColor(0x88,0x88,0x88))
        doc.add_paragraph(); return
    if not lines:
        p = doc.add_paragraph(_u("tx_no_history"))
        _set_font(p.runs[0] if p.runs else p.add_run(""), 10, color=RGBColor(0x88,0x88,0x88))
        doc.add_paragraph(); return
    if current:
        p = doc.add_paragraph()
        _set_font(p.add_run(_u("tx_current_status") + "  "), 9.5, bold=True, color=ACCENT)
        _set_font(p.add_run(current), 9.5, color=RGBColor(0x1A,0x1A,0x1A))
        p.paragraph_format.space_after = Pt(10)
    EFFICACY = {"CR": OK_T, "PR": OK_T, "SD": WARN_T, "PD": ALERT_T}
    LINE_NAMES = _u("tx_line_names").split(",")

    for idx, line in enumerate(lines):
        ah  = LINE_ACCENT_HEX[idx % len(LINE_ACCENT_HEX)]
        bh  = LINE_BG_HEX[idx % len(LINE_BG_HEX)]
        ar  = RGBColor(int(ah[0:2], 16), int(ah[2:4], 16), int(ah[4:6], 16))
        line_n = line.get("line", idx + 1)
        try:
            line_name = LINE_NAMES[int(line_n) - 1]
        except Exception:
            line_name = _u("tx_line_fallback").format(n=line_n)

        # ── 卡片：2列表（左色标 + 内容列）────────────────────────────
        BADGE_W = 680
        card = doc.add_table(rows=1, cols=2)
        card.alignment = WD_TABLE_ALIGNMENT.LEFT
        _fix_table(card, [BADGE_W, CONTENT_DXA - BADGE_W])
        card.rows[0].cells[0]._tc.get_or_add_tcPr()

        # 左侧色标列（纯色竖条）
        bc = card.rows[0].cells[0]
        _shading(bc, ah)
        _borders(bc, color=ah, size=4)
        bp = bc.paragraphs[0]
        bp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        bp.paragraph_format.space_before = Pt(8)
        bp.paragraph_format.space_after  = Pt(4)
        _set_font(bp.add_run(line_name), 9.5, bold=True, color=RGBColor(0xFF, 0xFF, 0xFF))

        # 右侧内容列
        cc = card.rows[0].cells[1]
        _shading(cc, bh)
        _borders(cc, color="E0E8F0", size=4)

        # 标题行：方案名 + 期间
        ph = cc.paragraphs[0]
        ph.paragraph_format.space_before = Pt(8)
        ph.paragraph_format.space_after  = Pt(3)
        ph.paragraph_format.left_indent  = Cm(0.4)
        _set_font(ph.add_run(line.get("regimen", "—")), 11, bold=True, color=PRIMARY)
        period = line.get("period", "")
        if period:
            _set_font(ph.add_run(f"  |  {period}"), 9, color=RGBColor(0x66, 0x66, 0x66))

        # 详细字段
        for lbl, key, use_ec in [(_u("tx_efficacy"), "efficacy", True),
                                  (_u("tx_stop_reason"), "stop_reason", False),
                                  (_u("tx_toxicity"), "toxicity", False)]:
            val = line.get(key, "")
            if not val:
                continue
            pd_p = cc.add_paragraph()
            pd_p.paragraph_format.left_indent  = Cm(0.4)
            pd_p.paragraph_format.space_before = Pt(2)
            pd_p.paragraph_format.space_after  = Pt(2)
            _set_font(pd_p.add_run(f"{lbl}{_u('kv_separator')}"), 9, bold=True, color=RGBColor(0x55, 0x55, 0x55))
            vc = RGBColor(0x1A, 0x1A, 0x1A)
            if use_ec:
                for k, col in EFFICACY.items():
                    if k in str(val):
                        vc = col
                        break
            _set_font(pd_p.add_run(str(val)), 9.5, color=vc)

        cc.add_paragraph().paragraph_format.space_after = Pt(6)
        doc.add_paragraph().paragraph_format.space_after = Pt(4)
    doc.add_paragraph()


# ── Gaps ───────────────────────────────────────────────────────────────────────
def gaps_section(doc, gaps):
    critical    = gaps.get("critical", [])
    recommended = gaps.get("recommended", [])
    covered     = gaps.get("covered", [])
    ACTION_LBL  = {"现医院补检": _u("action_cat_hospital"), "调阅历史档案": _u("action_cat_archive"),
                   "转诊专项检查": _u("action_cat_referral"), "组织已不可及": _u("action_cat_unavail")}
    def gap_item(g, tc):
        p = doc.add_paragraph()
        p.paragraph_format.left_indent  = Cm(0.5)
        p.paragraph_format.space_before = Pt(3)
        p.paragraph_format.space_after  = Pt(0)
        _set_font(p.add_run(f"- {g.get('item','')}"), 10, bold=True, color=tc)
        if g.get("reason"):
            _set_font(p.add_run(f"  --  {g['reason']}"), 9.5, color=RGBColor(0x55,0x55,0x55))
        if g.get("action_detail"):
            p2 = doc.add_paragraph()
            p2.paragraph_format.left_indent  = Cm(1.2)
            p2.paragraph_format.space_before = Pt(1)
            p2.paragraph_format.space_after  = Pt(4)
            lbl = ACTION_LBL.get(g.get("action_category",""), g.get("action_category",""))
            if lbl:
                _set_font(p2.add_run(f"[{lbl}]  "), 9, bold=True, color=RGBColor(0x44,0x44,0x44))
            _set_font(p2.add_run(g["action_detail"]), 9, color=RGBColor(0x44,0x44,0x44))
    if critical:
        p = doc.add_paragraph()
        _set_font(p.add_run(_u("gaps_critical_title")), 10, bold=True, color=ALERT_T)
        p.paragraph_format.space_before = Pt(6)
        for g in critical: gap_item(g, ALERT_T)
    if recommended:
        p = doc.add_paragraph()
        _set_font(p.add_run(_u("gaps_recommended_title")), 10, bold=True, color=WARN_T)
        p.paragraph_format.space_before = Pt(8)
        for g in recommended: gap_item(g, WARN_T)
    if covered:
        p = doc.add_paragraph()
        _set_font(p.add_run(_u("gaps_covered_title")), 10, bold=True, color=OK_T)
        p.paragraph_format.space_before = Pt(8)
        for g in covered:
            p2 = doc.add_paragraph()
            p2.paragraph_format.left_indent  = Cm(0.5)
            p2.paragraph_format.space_before = Pt(3)
            p2.paragraph_format.space_after  = Pt(2)
            _set_font(p2.add_run(f"- {g.get('item','')}"), 10, bold=True, color=OK_T)
            if g.get("reason"):
                _set_font(p2.add_run(f"  --  {g['reason']}"), 9.5, color=RGBColor(0x55,0x55,0x55))
    doc.add_paragraph()


# ── Review flags ───────────────────────────────────────────────────────────────
def review_flags_section(doc, flags):
    if not flags:
        p = doc.add_paragraph(_u("flags_all_clear"))
        _set_font(p.runs[0] if p.runs else p.add_run(""), 10, color=OK_T)
        return
    red    = [f for f in flags if f.get("severity")=="red"]
    yellow = [f for f in flags if f.get("severity")=="yellow"]
    if red:
        p = doc.add_paragraph()
        _set_font(p.add_run(_u("flags_red_title").format(n=len(red))), 10, bold=True, color=ALERT_T)
        p.paragraph_format.space_before = Pt(4)
        for f in red:
            p2 = doc.add_paragraph()
            p2.paragraph_format.left_indent  = Cm(0.5)
            p2.paragraph_format.space_before = Pt(2)
            _set_font(p2.add_run(f"[{f.get('id','RF')}]  {f.get('issue','')}"), 9.5, color=ALERT_T)
            if f.get("suggested_action"):
                p3 = doc.add_paragraph()
                p3.paragraph_format.left_indent = Cm(1.2)
                _set_font(p3.add_run(_u("flags_suggested_action").format(action=f['suggested_action'])), 9, color=RGBColor(0x55,0x55,0x55))
    if yellow:
        p = doc.add_paragraph()
        _set_font(p.add_run(_u("flags_yellow_title").format(n=len(yellow))), 10, bold=True, color=WARN_T)
        p.paragraph_format.space_before = Pt(6)
        for f in yellow:
            p2 = doc.add_paragraph()
            p2.paragraph_format.left_indent  = Cm(0.5)
            p2.paragraph_format.space_before = Pt(2)
            _set_font(p2.add_run(f"[{f.get('id','RF')}]  {f.get('issue','')}"), 9.5, color=WARN_T)
    doc.add_paragraph()


# ── Pathway (current treatment route) ─────────────────────────────────────────
def pathway_section(doc, pathway):
    """当前治疗路径 - 高亮文本块（参考 004 PDF 样式）。"""
    if not pathway:
        return
    # 外框卡片（浅灰蓝边框）
    card = doc.add_table(rows=1, cols=1)
    card.alignment = WD_TABLE_ALIGNMENT.LEFT
    _fix_table(card, [CONTENT_DXA])
    cc = card.rows[0].cells[0]
    _shading(cc, "F7FAFD")
    _borders(cc, color="C5D8EC", size=6)

    LABEL_COLORS = {
        "current":     (ACCENT,  _u("pw_current")),
        "bridge":      (WARN_T,  _u("pw_bridge")),
        "alternative": (PRIMARY, _u("pw_alternative")),
    }

    def _path_para(cell, label, text, lbl_color):
        p = cell.add_paragraph()
        p.paragraph_format.left_indent  = Cm(0.4)
        p.paragraph_format.space_before = Pt(3)
        p.paragraph_format.space_after  = Pt(3)
        _set_font(p.add_run(f"{label}{_u('kv_separator')}"), 9.5, bold=True, color=lbl_color)
        _set_font(p.add_run(str(text)), 9.5, color=RGBColor(0x1A, 0x1A, 0x1A))

    first = True
    for key, (col, label) in LABEL_COLORS.items():
        val = pathway.get(key, "")
        if not val:
            continue
        p = cc.paragraphs[0] if first else cc.add_paragraph()
        if first:
            p.paragraph_format.space_before = Pt(8)
            p.paragraph_format.space_after  = Pt(3)
            p.paragraph_format.left_indent  = Cm(0.4)
            _set_font(p.add_run(f"{label}{_u('kv_separator')}"), 9.5, bold=True, color=col)
            _set_font(p.add_run(str(val)), 9.5, color=RGBColor(0x1A, 0x1A, 0x1A))
            first = False
        else:
            p.paragraph_format.left_indent  = Cm(0.4)
            p.paragraph_format.space_before = Pt(3)
            p.paragraph_format.space_after  = Pt(3)
            _set_font(p.add_run(f"{label}{_u('kv_separator')}"), 9.5, bold=True, color=col)
            _set_font(p.add_run(str(val)), 9.5, color=RGBColor(0x1A, 0x1A, 0x1A))

    # pending_issues / next_steps fallback
    for issue in pathway.get("pending_issues", []):
        pi = cc.add_paragraph()
        pi.paragraph_format.left_indent  = Cm(0.4)
        pi.paragraph_format.space_before = Pt(2)
        pi.paragraph_format.space_after  = Pt(2)
        _set_font(pi.add_run(f"- {issue}"), 9.5, color=RGBColor(0x1A, 0x1A, 0x1A))
    ns = pathway.get("next_steps", "")
    if ns and not pathway.get("current"):
        pns = cc.add_paragraph()
        pns.paragraph_format.left_indent  = Cm(0.4)
        pns.paragraph_format.space_before = Pt(3)
        pns.paragraph_format.space_after  = Pt(6)
        _set_font(pns.add_run(str(ns)), 9.5, color=RGBColor(0x1A, 0x1A, 0x1A))
    else:
        cc.add_paragraph().paragraph_format.space_after = Pt(6)

    doc.add_paragraph()


# ── Sources ────────────────────────────────────────────────────────────────────
def sources_table(doc, sources):
    if not sources: return
    col_widths = [1633, 3266, 4173]
    tbl = doc.add_table(rows=1, cols=3)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    _fix_table(tbl, col_widths)
    for i, h in enumerate([_u("src_th_module"), _u("src_th_datapoint"), _u("src_th_file")]):
        c = tbl.rows[0].cells[i]; _shading(c, HDR_BG); _borders(c)
        _set_font(c.paragraphs[0].add_run(h), 9.5, bold=True, color=PRIMARY)
    for s in sources:
        row = tbl.add_row()
        for i, val in enumerate([s.get("module",""), s.get("field",""), s.get("file","")]):
            c = row.cells[i]; _borders(c)
            p = c.paragraphs[0]
            p.paragraph_format.space_before = Pt(2); p.paragraph_format.space_after = Pt(2)
            _set_font(p.add_run(str(val)), 9)
    doc.add_paragraph()


# ── Disclaimer ─────────────────────────────────────────────────────────────────
def disclaimer(doc, data):
    doc.add_paragraph()
    rule = doc.add_paragraph(); _bottom_rule(rule, color="CCCCCC", size=6)
    rule.paragraph_format.space_after = Pt(4)
    gen   = data.get("generated_at", datetime.now().isoformat())
    fn    = data.get("files_analyzed", 0)
    fr, fy, fg = data.get("review_flags_red",0), data.get("review_flags_yellow",0), data.get("review_flags_green",0)
    tmpl = _u("disclaimer")
    txt  = tmpl.format(gen=gen[:19], fn=fn, fr=fr, fy=fy, fg=fg)
    p2 = doc.add_paragraph()
    _set_font(p2.add_run(txt), 8, color=RGBColor(0x88,0x88,0x88))
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER


def _kv(doc, label, val):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(1); p.paragraph_format.space_after = Pt(1)
    _set_font(p.add_run(f"{label}{_u('kv_separator')}"), 10, bold=True, color=RGBColor(0x44,0x44,0x44))
    pending = any(kw in str(val) for kw in _u("kv_pending_keywords").split(","))
    _set_font(p.add_run(str(val)), 10,
              color=RGBColor(0x88,0x88,0x88) if pending else RGBColor(0x1A,0x1A,0x1A))


# ── Builders ───────────────────────────────────────────────────────────────────
def build_brief(data, output_path):
    global _UI
    _UI = data.get("ui", {})
    doc = Document(); _page_setup(doc); _footer(doc)
    cover_brief(doc, data)
    patient = data.get("patient", {})
    dx      = data.get("diagnosis", {})
    heading_brief(doc, _u("h_patient_id"))
    patient_info_table(doc, patient, extended=False)
    heading_brief(doc, _u("h_condition"))
    for lbl, key in [(_u("kv_date"),"date"),(_u("kv_site"),"primary_site"),(_u("kv_hist"),"histology"),
                     (_u("kv_stage"),"stage"),(_u("kv_mets"),"metastasis"),
                     (_u("kv_init"),"initial_or_recurrence"),(_u("kv_stat"),"current_status")]:
        _kv(doc, lbl, dx.get(key,"—"))
    doc.add_paragraph()
    km = [m for m in data.get("molecular",[]) if m.get("priority")=="high"]
    if km:
        heading_brief(doc, _u("h_molecular"))
        molecular_table(doc, km)
    img = data.get("imaging",{})
    if img and (img.get("items") or img.get("note")):
        heading_brief(doc, _u("h_imaging"))
        imaging_section(doc, img)
    heading_brief(doc, _u("h_labs"))
    labs = data.get("labs", [])
    tkeys = _trend_keys(labs)
    labs_trend_charts(doc, labs, data.get("trend_events", []))
    labs_snapshot_chart(doc, labs)
    latest_labs = _latest_per_trend(labs, tkeys)
    if latest_labs:
        labs_cards(doc, latest_labs)
    tx = data.get("treatment",{})
    if tx.get("lines") or tx.get("note"):
        heading_brief(doc, _u("h_treatment"))
        treatment_history(doc, tx)
    gp = data.get("gaps",{})
    if gp.get("critical") or gp.get("recommended") or gp.get("covered"):
        heading_brief(doc, _u("h_gaps"))
        gaps_section(doc, gp)
    pathway = data.get("pathway", {})
    if pathway:
        heading_brief(doc, _u("h_pathway"))
        pathway_section(doc, pathway)
    flags = [f for f in data.get("review_flags",[]) if not f.get("user_confirmed")]
    if flags:
        heading_brief(doc, _u("h_flags"))
        review_flags_section(doc, flags)
    disclaimer(doc, data)
    doc.save(output_path)
    print(f"[brief] Saved: {output_path}")

def build_detailed(data, output_path):
    global _UI
    _UI = data.get("ui", {})
    doc = Document(); _page_setup(doc); _footer(doc)
    cover_detailed(doc, data)
    patient = data.get("patient", {})
    dx      = data.get("diagnosis", {})
    heading_numbered(doc, 1, _u("s1"))
    patient_info_table(doc, patient, extended=True)
    heading_numbered(doc, 2, _u("s2"))
    for lbl, key in [(_u("kv_date"),"date"),(_u("kv_site"),"primary_site"),(_u("kv_hist"),"histology"),
                     (_u("kv_diff"),"differentiation"),(_u("kv_stage"),"stage"),
                     (_u("kv_init"),"initial_or_recurrence"),(_u("kv_mets"),"metastasis"),
                     (_u("kv_stat"),"current_status")]:
        _kv(doc, lbl, dx.get(key,"—"))
    doc.add_paragraph()
    heading_numbered(doc, 3, _u("s3"))
    molecular_table(doc, data.get("molecular",[]))
    heading_numbered(doc, 4, _u("s4"))
    imaging_section(doc, data.get("imaging",{}))
    heading_numbered(doc, 5, _u("s5"))
    labs_table(doc, data.get("labs", []))
    labs_trend_charts(doc, data.get("labs", []), data.get("trend_events", []))
    labs_snapshot_chart(doc, data.get("labs", []))
    heading_numbered(doc, 6, _u("s6"))
    treatment_history(doc, data.get("treatment",{}))
    pathway = data.get("pathway", {})
    if pathway:
        heading_numbered(doc, 7, _u("s7"))
        pathway_section(doc, pathway)
    heading_numbered(doc, "A", _u("sA"))
    gaps_section(doc, data.get("gaps",{}))
    heading_numbered(doc, "B", _u("sB"))
    review_flags_section(doc, [f for f in data.get("review_flags",[]) if not f.get("user_confirmed")])
    if data.get("sources"):
        heading_numbered(doc, "C", _u("sC"))
        sources_table(doc, data["sources"])
    disclaimer(doc, data)
    doc.save(output_path)
    print(f"[detailed] Saved: {output_path}")

# ── Main ───────────────────────────────────────────────────────────────────────
def main():
    parser = argparse.ArgumentParser()
    parser.add_argument("data_json");  parser.add_argument("output_docx")
    parser.add_argument("--type", choices=["brief","detailed"], default="brief")
    parser.add_argument("--md-patch", metavar="MD_FILE",
                        help="同时将快照图保存为 PNG 并插入指定 .md 文件")
    args = parser.parse_args()
    with open(args.data_json, "r", encoding="utf-8") as f:
        data = json.load(f)
    (build_brief if args.type=="brief" else build_detailed)(data, args.output_docx)

    if args.md_patch:
        import os, re as _re
        md_path = args.md_patch
        md_dir  = os.path.dirname(os.path.abspath(md_path))
        charts_dir = os.path.join(md_dir, "charts")
        png_name   = f"labs_snapshot_{args.type}.png"
        png_path   = os.path.join(charts_dir, png_name)
        saved = labs_snapshot_chart(None, data.get("labs", []), png_save_path=png_path)
        if saved and os.path.exists(md_path):
            with open(md_path, "r", encoding="utf-8") as f:
                content = f.read()
            alt_text = _u("chart_snapshot_title")
            img_tag = f"\n\n![{alt_text}](./charts/{png_name})\n"
            patterns = [
                r"(##\s*模块\s*5[^\n]*\n)",
                r"(##\s*关键实验室指标[^\n]*\n)",
            ]
            patched = False
            for pat in patterns:
                m = _re.search(pat, content)
                if m:
                    section_start = m.end()
                    next_heading = _re.search(r"\n##\s+", content[section_start:])
                    section_end  = section_start + next_heading.start() if next_heading else len(content)
                    if img_tag.strip() not in content:
                        content = content[:section_end] + img_tag + content[section_end:]
                    patched = True
                    break
            if not patched and img_tag.strip() not in content:
                content += img_tag
            with open(md_path, "w", encoding="utf-8") as f:
                f.write(content)
            print(f"[md-patch] PNG -> {png_path}")
            print(f"[md-patch] Patched: {md_path}")

if __name__ == "__main__":
    main()
